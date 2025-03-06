import tkinter as tk
import os
import subprocess
import datetime as dt

from tkinter import filedialog
from tkinter import messagebox

import docx


padding_options={'fill': 'x', 'expand': True, 'padx': 5, 'pady': 2}



class InvoiceAutomation:
	def __init__(self):
		self.window=tk.Tk()
		self.window.geometry("500x600")
		self.window.title("Invoice automation")
		self.payment_methods=self.initialize_payments()
		self.create_widgets()
		self.create_invoice()
		
	def initialize_payments(self):
		return {
			'main bank':{
				'Recipient':'Gerald & Co',
				'Bank':'Absa Bank',
				'Account No':'123456789',
				'BIC':'ABCDEFG'

			},
			'Second bank':{
				'Recipient':'Gerald & Co',
				'Bank':'Family Bank',
				'Account No':'123456789',
				'BIC':'ABCDEFG'

			},
			'Private bank':{
				'Recipient':'Gerald & Co',
				'Bank':'Equity Bank',
				'Account No':'123456789',
				'BIC':'ABCDEFG'

			}
		}
	def create_widgets(self):
		self.create_invoice_labels_and_entries()
		
		self.create_payment_method()
		self.create_invoice_button()


	def create_invoice_labels_and_entries(self):
		self.partner_label=tk.Label(self.window, text="Partner")
		self.partner_label.pack(padding_options)
		self.partner_entry=tk.Entry(self.window)
		self.partner_entry.pack(padding_options)
		self.partner_street_label=tk.Label(self.window, text="Partner_street")
		self.partner_street_label.pack(padding_options)
		self.partner_street_entry=tk.Entry(self.window)
		self.partner_street_entry.pack(padding_options)
		self.partner_zip_city_country_label=tk.Label(self.window, text="Partner_zip_city_country")
		self.partner_zip_city_country_label.pack(padding_options)
		self.partner_zip_city_country_entry=tk.Entry(self.window)
		self.partner_zip_city_country_entry.pack(padding_options)
		self.invoice_number_label=tk.Label(self.window, text="Invoice_number")
		self.invoice_number_label.pack(padding_options)
		self.invoice_number_entry=tk.Entry(self.window)
		self.invoice_number_entry.pack(padding_options)
		self.service_description_label=tk.Label(self.window, text="service_description")
		self.service_description_label.pack(padding_options)
		self.service_description_entry=tk.Entry(self.window)
		self.service_description_entry.pack(padding_options)
		self.service_amount_label=tk.Label(self.window, text="service_amount")
		self.service_amount_label.pack(padding_options)
		self.service_amount_entry=tk.Entry(self.window)
		self.service_amount_entry.pack(padding_options)
		self.service_single_price_label=tk.Label(self.window, text="service_single_price")
		self.service_single_price_label.pack(padding_options)
		self.service_single_price_entry=tk.Entry(self.window)
		self.service_single_price_entry.pack(padding_options)
		
		

	
		
		
		
		
		
		
		

	


	def create_payment_method(self):
		self.payment_method=tk.StringVar(self.window)
		self.payment_method.set('main bank')
		self.payment_method_dropdown=tk.OptionMenu(self.window, self.payment_method,"main bank", "Family bank","Private bank")
		
		self.payment_method_label=tk.Label(self.window, text="payment_method")
		self.payment_method_label.pack(padding_options)
		self.payment_method_entry=tk.Entry(self.window)
		self.payment_method_entry.pack(padding_options)
		
		

		self.payment_method_dropdown.pack(padding_options)


	def create_invoice_button(self):
		self.create_button=tk.Button(self.window, text="create invoice", command=self.create_invoice)
		self.create_button.pack(padding_options)
		







	@staticmethod
	def replace_text(paragraph, old_text, new_text):
		if old_text in paragraph.text:
			paragraph.text=paragraph.text.replace(old_text, new_text)



	def create_invoice(self):
		doc=docx.Document('invoice_template.docx')	

		selected_payments_method=self.payment_methods[self.payment_method.get()]

		try:
			replacements={
				"[date]":dt.datetime.today().strftime('%y-%m-%/d'), # type: ignore
				"[partner]":self.partner_entry.get(),#type:ignore
				"[partner street]":self.partner_street_entry.get(), # type: ignore
				"[partner zip city country]":self.partner_zip_city_country_entry.get(),# type:ignore

				"[invoice number]":self.invoice_number_entry.get(),# type:ignore
				"[service description]":self.service_description_entry.get(),# type:ignore
				"[amount]":self.amount_entry.get(),#type: ignore

				"[single price]":f"${float(self.service_single_price_entry.get()):.2f}",# type:ignore
				"[full price]":f"${float(self.amount_entry.get()) * float(self.service_single_price_entry.get()):.2f}",# type:ignore
				"[Recipient]":selected_payment_method['Recipient'], #type:ignore
				"[Bank]":selected_payment_method['Bank'],# type:ignore
				"[Account No]":selected_payment_method['Account No'],# type:ignore
				"[BIC]":selected_payment_method['BIC']# type:ignore
					
			} # type: ignore
		
			for paragraph in list(doc.paragraphs):
				for old_text, new_text in replacements.items():
					self.replace_text(paragraph, old_text, new_text)
			for table in doc.tables:
				for row in table.rows:
					for  cell in row.cells:
						for paragraphs in cell.paragraphs:
							for old_text, new_text in replacements.item():
								self.replace_text(paragraph, old_text, new_text)

		except ValueError:	
			messagebox.showerror('invalid amount or price')
					

		save_path=filedialog.asksaveasfilename(defaultextension='.pdf',filetypes=[('PDF documents', '*.pdf')])
		doc.save('filled.docx')

		subprocess.run(['libreoffice','--headless','--convert-to','pdf','filled.docx','--outdir','.'], check=True)
		os.rename('filled.pdf', save_path) # type: ignore


		messagebox.showinfo('invoice created and saved successfullly')


			#docx2pdf
			#pdf2docx


		
if __name__=="__main__":
	InvoiceAutomation()
				

